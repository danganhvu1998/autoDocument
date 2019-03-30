import openpyxl
import docx
import re
import autoFillController

folderName = ""
files = []
students = {}
translates = {}
translates[""] = "[[[XXX- NO TRANSLATION FOUND -XXX]]]"

#Both docx and xlsx
def autoLine(paraText):
    global students, translates
    # JAPANESE - NEED TRANSLATION
    paraForms = re.findall(r"\[\[\[([a-zA-Z0-9_\.\-]+)\.nihon\]\]\]", paraText)
    for paraForm in paraForms:
        form = "[[["+paraForm+".nihon]]]"
        #data = students.get(paraForm, "")
        data = translates[students.get(paraForm, "").lower()]
        print(1, paraForm, form, data)
        paraText = paraText.replace(form, data)

    # MIGHT HAVE SOME ADDITIONAL CONDITION
    paraForms = re.findall(r"\[\[\[([a-zA-Z0-9_\.\-]+)\]\]\]", paraText)
    for paraForm in paraForms: #paraFrom is like name.first.english.capital
        form = "[[["+paraForm+"]]]"
        elements = paraForm.split(".")
        rawData = students.get(elements[0], "")
        elements = elements[1::]
        data = autoFillController.autoForm(rawData, elements)
        paraText = paraText.replace(form, data)
    return paraText

#DOCX
def autoParagraph(paragraph):
    lines = paragraph.runs #split to many same style parts
    for line in lines: 
        line.text = autoLine(line.text)
    return paragraph

def autoTable(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph = autoParagraph(paragraph)
            for cellTable in cell.tables:
                autoTable(cellTable)
    return table

def docxAutoDocument(file):
    global folderName
    fileUrl = 'temp/'+folderName+'/'+file
    document = docx.Document(docx = fileUrl)

    # PARAGRAPHS
    for paragraph in document.paragraphs:
        paragraph = autoParagraph(paragraph)
    
    # TABLES
    for table in document.tables:
        table = autoTable(table)
                        
    document.save(fileUrl)
    return 1

#XLSX
def checkNullRow(ws, row): #return 0: row is null, 1: row is not null
    for col in range(1,21):
        if( ws.cell(column=col, row=row).value != None ): return 0
    return 1

def xlsxAutoDocument(file):
    #Open File
    global folderName
    fileUrl = 'temp/'+folderName+'/'+file
    wb = openpyxl.load_workbook(filename = fileUrl)
    #Set Null Accept Distance
    nullAcceptDistance = 20


    #Rock N Roll!
    for ws in wb:
        currNullRow = 0
        currRow = 1
        while(1):
            #print(ws, currRow)
            if(currNullRow>nullAcceptDistance): break
            if(checkNullRow(ws, currRow)):
                currNullRow += 1
                currRow += 1
                continue
            else:
                currCol = 1
                currNullRow = 0
                currNullCol = 0
                while(1):
                    if(currNullCol>nullAcceptDistance): break
                    if( ws.cell(column=currCol, row=currRow).value == None ):
                        currNullCol += 1
                        currCol += 1
                        continue
                    else:
                        currNullCol = 0
                        ws.cell(column=currCol, row=currRow).value = autoLine( ws.cell(column=currCol, row=currRow).value)
                        currCol += 1
                currRow += 1
    wb.save(fileUrl)

#MAIN
def autoDocument(file):
    if(file.endswith(".docx")):
        return docxAutoDocument(file)
    else:
        return xlsxAutoDocument(file)

def MAIN(clientFolderName, clientFiles, clientStudents, clientTranslates):
    #PRERUN ASSIGN VALUE
    global folderName, files, students, translates
    folderName = clientFolderName
    files = clientFiles
    for clientStudent in clientStudents:
        students[clientStudent[1]] = clientStudent[0]
        if( clientStudent[2] != None  ): students[clientStudent[2]] = clientStudent[0]
    
    for clientTranslate in clientTranslates:
        translates[clientTranslate[0].lower()] = clientTranslate[1]
    
    #AUTO DOCUMENT
    for file in files:
        autoDocument(file)

#cd temp/Toyo_University___Đặng_Anh_Vũ/ && mv table.docx table.zip && 
